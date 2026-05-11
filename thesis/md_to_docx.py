import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DEFAULT_MD_PATH = ROOT / "复杂遮挡环境下飞行器连续稳定定位技术_初稿.md"
DEFAULT_DOCX_PATH = ROOT / "复杂遮挡环境下飞行器连续稳定定位技术_初稿.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=None):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_doc_defaults(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True

    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_markdown_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) >= 2 and all(set(c.replace(":", "").replace("-", "")) == set() for c in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=max(len(r) for r in parsed))
    table.style = "Table Grid"
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=9, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_paragraph_with_inline_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    parts = text.split("`")
    for i, part in enumerate(parts):
        run = p.add_run(part)
        set_run_font(run, size=11)
        if i % 2 == 1:
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def add_image(doc, line, base_dir):
    alt_start = line.find("[")
    alt_end = line.find("]")
    path_start = line.find("(", alt_end)
    path_end = line.find(")", path_start)
    caption = line[alt_start + 1 : alt_end]
    rel_path = line[path_start + 1 : path_end]
    image_path = base_dir / rel_path
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        set_run_font(run, size=10)


def convert(md_path=DEFAULT_MD_PATH, docx_path=DEFAULT_DOCX_PATH):
    doc = Document()
    set_doc_defaults(doc)
    md_path = Path(md_path)
    docx_path = Path(docx_path)
    base_dir = md_path.resolve().parent
    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(9)
                in_code = False
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].rstrip())
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        if line.startswith("!["):
            add_image(doc, line, base_dir)
            i += 1
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            if not doc.paragraphs or not doc.paragraphs[-1].text:
                pass
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("$$"):
            formula = []
            i += 1
            while i < len(lines) and not lines[i].startswith("$$"):
                formula.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("\n".join(formula))
            set_run_font(run, size=10)
            i += 1
            continue

        if line.startswith("- ") or line.startswith("（"):
            p = doc.add_paragraph(line, style=None)
            p.paragraph_format.first_line_indent = Pt(0)
            for run in p.runs:
                set_run_font(run, size=11)
            i += 1
            continue

        add_paragraph_with_inline_code(doc, line)
        i += 1

    doc.save(docx_path)
    print(docx_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", default=str(DEFAULT_MD_PATH))
    parser.add_argument("--output", default=str(DEFAULT_DOCX_PATH))
    args = parser.parse_args()
    convert(args.input, args.output)
